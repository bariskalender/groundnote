"""DOCX parser backed by python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document as open_docx_document
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from groundnote.documents.errors import CorruptDocumentError
from groundnote.documents.models import ParsedDocument, ParsedSection
from groundnote.documents.parsers.base import make_section, require_sections
from groundnote.domain import SupportedFileType


class DocxParser:
    """Extract paragraphs, headings, and simple table text from DOCX files."""

    supported_file_type = SupportedFileType.DOCX

    def validate_compatible(self, file_path: Path) -> None:
        self._load(file_path)

    def parse(
        self,
        file_path: Path,
        *,
        original_filename: str,
        stored_filename: str,
        sha256: str,
        file_size_bytes: int,
    ) -> ParsedDocument:
        document = self._load(file_path)
        sections: list[ParsedSection] = []
        warnings: list[str] = []
        current_heading: str | None = None
        buffer: list[str] = []
        order = 0

        def flush() -> None:
            nonlocal buffer, order
            if not buffer:
                return
            section = make_section(
                "\n".join(buffer),
                source_order=order,
                section_title=current_heading,
            )
            if section is not None:
                sections.append(section)
                order += 1
            buffer = []

        for paragraph in document.paragraphs:
            text = paragraph.text
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading") and text.strip():
                flush()
                current_heading = text.strip()
                buffer.append(text)
            elif text.strip():
                buffer.append(text)
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                buffer.append("\n".join(rows))
        flush()
        require_sections(sections)
        return ParsedDocument(
            original_filename=original_filename,
            stored_filename=stored_filename,
            file_type=self.supported_file_type,
            sha256=sha256,
            file_size_bytes=file_size_bytes,
            page_count=None,
            sections=sections,
            warnings=warnings,
        )

    @staticmethod
    def _load(file_path: Path) -> DocxDocument:
        try:
            return open_docx_document(str(file_path))
        except (PackageNotFoundError, OSError, ValueError) as exc:
            raise CorruptDocumentError() from exc
