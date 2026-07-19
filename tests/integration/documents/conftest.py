from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


@pytest.fixture
def document_dir(tmp_path: Path) -> Path:
    directory = tmp_path / "documents"
    directory.mkdir()
    return directory


def write_text_pdf(path: Path, pages: list[str]) -> Path:
    writer = PdfWriter()
    for text in pages:
        page = writer.add_blank_page(width=300, height=300)
        if text:
            font = DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
            font_ref = writer._add_object(font)  # noqa: SLF001 - pypdf test fixture helper.
            page[NameObject("/Resources")] = DictionaryObject(
                {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
            )
            stream = DecodedStreamObject()
            safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            stream.set_data(f"BT /F1 12 Tf 50 250 Td ({safe_text}) Tj ET".encode())
            page[NameObject("/Contents")] = writer._add_object(stream)  # noqa: SLF001
    with path.open("wb") as file:
        writer.write(file)
    return path


def write_encrypted_pdf(path: Path) -> Path:
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    writer.encrypt("fixture-passphrase")
    with path.open("wb") as file:
        writer.write(file)
    return path


def write_docx(path: Path) -> Path:
    document = DocxDocument()
    document.add_heading("Lecture Overview", level=1)
    document.add_paragraph("First paragraph with Türkçe characters and ∑ symbols.")
    document.add_paragraph("List item one", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Meaning"
    table.cell(1, 0).text = "RAG"
    table.cell(1, 1).text = "Retrieval augmented generation"
    document.save(path)
    return path


def write_empty_docx(path: Path) -> Path:
    DocxDocument().save(path)
    return path
